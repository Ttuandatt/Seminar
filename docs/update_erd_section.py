"""
Update section '1.1. Sơ đồ ERD đầy đủ 11 entity' with detailed entity descriptions.
Run: python -X utf8 docs/update_erd_section.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

SRC = os.path.join(os.path.dirname(__file__), 'Báo cáo Seminar - filled2.docx')
DST = SRC

doc = Document(SRC)

# ── Find the heading and its content paragraphs ──
heading_idx = None
for i, p in enumerate(doc.paragraphs):
    if '1.1. Sơ đồ ERD' in p.text:
        heading_idx = i
        break

if heading_idx is None:
    print('ERROR: heading not found')
    exit(1)

# Find the next heading to know where this section ends
next_heading_idx = None
for i in range(heading_idx + 1, len(doc.paragraphs)):
    style = doc.paragraphs[i].style
    if style and style.name and style.name.startswith('Heading'):
        next_heading_idx = i
        break

print(f'  Section 1.1 ERD: paragraphs {heading_idx} to {next_heading_idx - 1}')

# Remove old content paragraphs (between heading and next heading)
# We delete from the XML tree
elements_to_remove = []
for i in range(heading_idx + 1, next_heading_idx):
    elements_to_remove.append(doc.paragraphs[i]._element)

for el in elements_to_remove:
    el.getparent().remove(el)

print(f'  Removed {len(elements_to_remove)} old paragraphs')

# ── Helpers ──
def make_para(text, bold=False, size=10, indent_cm=0, space_after=3):
    """Create a paragraph element."""
    new_p = doc.element.makeelement(qn('w:p'), {})
    pPr = new_p.makeelement(qn('w:pPr'), {})

    # Indentation
    if indent_cm > 0:
        ind = new_p.makeelement(qn('w:ind'), {qn('w:left'): str(int(indent_cm * 567))})
        pPr.append(ind)

    # Spacing
    spacing = new_p.makeelement(qn('w:spacing'), {
        qn('w:after'): str(int(space_after * 20)),
        qn('w:line'): '276',
        qn('w:lineRule'): 'auto',
    })
    pPr.append(spacing)

    new_p.insert(0, pPr)

    # If text contains bold markers **...**
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        r = new_p.makeelement(qn('w:r'), {})
        rPr = new_p.makeelement(qn('w:rPr'), {})
        sz = new_p.makeelement(qn('w:sz'), {qn('w:val'): str(size * 2)})
        szCs = new_p.makeelement(qn('w:szCs'), {qn('w:val'): str(size * 2)})
        rPr.append(sz)
        rPr.append(szCs)

        is_bold = part.startswith('**') and part.endswith('**')
        if is_bold or bold:
            b_el = new_p.makeelement(qn('w:b'), {})
            rPr.append(b_el)
            part = part.strip('*')

        r.append(rPr)
        t = new_p.makeelement(qn('w:t'), {})
        t.text = part
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        new_p.append(r)

    return new_p

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex,
    })
    shading.append(shd)

def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)

def set_table_borders(tbl):
    tblPr = tbl._element.tblPr if tbl._element.tblPr is not None else tbl._element.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single', qn('w:sz'): '4',
            qn('w:space'): '0', qn('w:color'): 'B0B0B0',
        })
        borders.append(el)
    tblPr.append(borders)
    if tbl._element.tblPr is None:
        tbl._element.insert(0, tblPr)

# ── Build new content ──
ref = doc.paragraphs[heading_idx]._element

# We insert in REVERSE order because each addnext goes right after ref

# ── Entity summary table data ──
entities = [
    ('User', '13', '1:1 → ShopOwner\n1:1 → TouristUser\n1:N → Poi (creator)\n1:N → Poi (owner)\n1:N → Tour',
     'Bảng hợp nhất cho cả 3 vai trò (Admin, Shop Owner, Tourist). Phân quyền bằng enum Role thay vì tách 3 bảng riêng, đơn giản hóa authentication flow — chỉ cần 1 JWT strategy và 1 login endpoint cho tất cả.'),

    ('ShopOwner', '9', '1:1 ← User',
     'Tách thành extension table (1:1 với User) để lưu dữ liệu đặc thù chủ quán (shopName, phone, openingHours) mà không gây nullable columns trên bảng User chung. Chỉ tạo khi role = SHOP_OWNER.'),

    ('TouristUser', '10', '1:1 ← User\n1:N → Favorite\n1:N → ViewHistory',
     'Tương tự ShopOwner, tách riêng để lưu preferences của du khách (languagePref, autoPlayAudio, pushToken). Là FK owner cho Favorite và ViewHistory.'),

    ('Poi', '16', 'N:1 ← User (creator)\nN:1 ← User (owner)\n1:N → PoiMedia\nN:M ↔ Tour (qua TourPoi)\n1:N → Favorite\n1:N → ViewHistory\n1:N → TriggerLog',
     'Entity trung tâm của hệ thống. Chứa nội dung song ngữ (nameVi/nameEn, descriptionVi/descriptionEn), tọa độ GPS, triggerRadius, 8 loại category. Hỗ trợ soft delete (deletedAt) và 3 trạng thái (DRAFT → ACTIVE → ARCHIVED). Có 2 FK đến User: createdById (Admin tạo) và ownerId (Shop Owner sở hữu, nullable).'),

    ('PoiMedia', '13', 'N:1 ← Poi (CASCADE)',
     'Lưu file media đính kèm POI (IMAGE hoặc AUDIO). Tách riêng thay vì JSON array trên Poi để hỗ trợ query theo type/language, sắp xếp orderIndex, và cascade delete khi xóa POI. TTS audio mới nhất luôn có orderIndex = 0.'),

    ('Tour', '12', 'N:1 ← User (creator)\nN:M ↔ Poi (qua TourPoi)',
     'Tour tham quan gom nhiều POI theo lộ trình. Cùng cấu trúc song ngữ và soft delete như Poi. estimatedDuration giúp du khách ước tính thời gian.'),

    ('TourPoi', '4', 'N:1 ← Tour (CASCADE)\nN:1 ← Poi',
     'Junction table cho quan hệ N:M giữa Tour và Poi. orderIndex xác định thứ tự điểm dừng trong tour. UNIQUE(tourId, poiId) ngăn duplicate. CASCADE delete khi xóa Tour.'),

    ('Favorite', '4', 'N:1 ← TouristUser (CASCADE)\nN:1 ← Poi',
     'Lưu danh sách POI yêu thích của du khách. UNIQUE(touristId, poiId) đảm bảo mỗi du khách chỉ yêu thích 1 POI 1 lần. Tách riêng thay vì JSON array để query nhanh và đếm aggregation.'),

    ('ViewHistory', '6', 'N:1 ← TouristUser (CASCADE)\nN:1 ← Poi',
     'Ghi lịch sử xem POI cho du khách ĐÃ ĐĂNG NHẬP. Lưu triggerType (GPS/QR/MANUAL) và audioPlayed để phân tích hành vi. Tách khỏi TriggerLog vì chỉ áp dụng cho authenticated users.'),

    ('TriggerLog', '9', 'N:1 ← Poi',
     'Ghi nhật ký MỌI trigger event kể cả du khách CHƯA đăng nhập (identify bằng deviceId). Lưu tọa độ thực tế (userLat/userLng) và khoảng cách (distanceMeters) để phân tích độ chính xác GPS. Tách khỏi ViewHistory vì phục vụ mục đích analytics khác.'),

    ('PasswordResetToken', '6', 'N:1 ← User (CASCADE)',
     'Token đặt lại mật khẩu dùng 1 lần (usedAt != null = đã sử dụng), hết hạn sau 1 giờ. Tách bảng riêng thay vì lưu trên User để hỗ trợ nhiều request reset cùng lúc và audit trail.'),

    ('RevokedToken', '7', 'N:1 ← User (CASCADE)',
     'Lưu JWT token đã bị thu hồi (khi logout hoặc force-revoke). JwtAuthGuard kiểm tra bảng này trên mỗi request. expiresAt cho phép cleanup tự động token hết hạn. Cần thiết vì JWT là stateless — không có cách "invalidate" token nếu không track server-side.'),
]

# ── Create table ──
num_rows = 1 + len(entities)
table = doc.add_table(rows=num_rows, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table)

# Column widths
col_widths = [Cm(2.5), Cm(1.2), Cm(3.5), Cm(7.8)]
for row in table.rows:
    for j, w in enumerate(col_widths):
        row.cells[j].width = w

# Header
headers = ['Entity', 'Fields', 'Quan hệ', 'Mục đích & Lý do thiết kế']
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    set_cell_text(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=(255, 255, 255))
    set_cell_shading(cell, '0C4A6E')

# Data rows
for i, (name, fields, rels, purpose) in enumerate(entities):
    row = table.rows[i + 1]
    set_cell_text(row.cells[0], name, bold=True, size=9)
    set_cell_text(row.cells[1], fields, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    set_cell_text(row.cells[2], rels, size=8)
    set_cell_text(row.cells[3], purpose, size=9)
    if i % 2 == 1:
        for j in range(4):
            set_cell_shading(row.cells[j], 'F8FAFC')

# ── Move table from end-of-doc to correct position ──
table_el = table._element
table_el.getparent().remove(table_el)

# ── Build paragraphs to insert (in reverse order) ──
elements_to_insert = []

# 1. Intro paragraph
elements_to_insert.append(make_para(
    'Cơ sở dữ liệu gồm **12 entity** với tổng cộng **109 fields**, **10 enums**, và **20 indexes** (bao gồm 4 composite indexes). '
    'Dưới đây là bảng tổng hợp từng entity, mối quan hệ, mục đích và lý do thiết kế:',
    size=10, space_after=6
))

# 2. Table
elements_to_insert.append(table_el)

# 3. Design principles paragraph
elements_to_insert.append(make_para(
    '**Nguyên tắc thiết kế chính:**',
    size=10, bold=True, space_after=2
))

principles = [
    '**Unified User table:** 1 bảng User duy nhất cho 3 vai trò, phân quyền bằng enum Role. Extension tables (ShopOwner, TouristUser) quan hệ 1:1 chứa dữ liệu đặc thù — tránh nullable columns không cần thiết.',
    '**Tách ViewHistory và TriggerLog:** ViewHistory chỉ cho du khách đã đăng nhập (gắn touristId), TriggerLog ghi MỌI trigger kể cả anonymous (gắn deviceId). Phục vụ 2 mục đích phân tích khác nhau: hành vi cá nhân vs. analytics toàn hệ thống.',
    '**Soft delete cho Poi và Tour:** Dùng deletedAt thay vì DELETE FROM. Bảo toàn tham chiếu từ ViewHistory/TriggerLog/Favorite. User dùng status LOCKED/INACTIVE thay thế.',
    '**UUID v4 primary keys:** An toàn hơn auto-increment (không đoán được), phù hợp nếu mở rộng distributed system.',
    '**Bilingual inline fields:** nameVi/nameEn, descriptionVi/descriptionEn trực tiếp trên entity thay vì bảng translation riêng — đơn giản cho MVP 2 ngôn ngữ, tránh JOIN overhead.',
    '**Junction table TourPoi:** Cho phép 1 POI thuộc nhiều Tour và thứ tự tùy chỉnh (orderIndex), khác với embedded array không hỗ trợ reorder.',
]

for pr in principles:
    elements_to_insert.append(make_para(f'• {pr}', size=9, indent_cm=0.5, space_after=3))

# ── Insert all elements after heading (reverse order) ──
for el in reversed(elements_to_insert):
    ref.addnext(el)

print(f'  + Inserted intro + table (12 entities) + 6 design principles')

doc.save(DST)
print(f'\nSaved to {DST}')
