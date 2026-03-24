"""
Add RBAC permission matrix table to the filled report.
Run: python -X utf8 docs/add_rbac_table.py
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

SRC = os.path.join(os.path.dirname(__file__), 'Báo cáo Seminar - filled.docx')
DST = os.path.join(os.path.dirname(__file__), 'Báo cáo Seminar - filled2.docx')

doc = Document(SRC)

# ── Find the heading ──
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if '3.2. Ma trận phân quyền' in p.text:
        target_idx = i
        break

if target_idx is None:
    print('ERROR: heading "3.2. Ma trận phân quyền" not found')
    exit(1)

# ── Find the insertion point: after the last paragraph of section 3.2 ──
# (the content paragraphs inserted by fill_report.py come AFTER the heading)
# We need to find the element after the existing content paragraphs
# Look for the next heading or section break
insert_after_idx = target_idx
for i in range(target_idx + 1, len(doc.paragraphs)):
    p = doc.paragraphs[i]
    style = p.style
    if style and style.name and style.name.startswith('Heading'):
        break
    if p.text.strip():  # non-empty paragraph in this section
        insert_after_idx = i

print(f'  Found heading at paragraph {target_idx}')
print(f'  Will insert table after paragraph {insert_after_idx}')

# ── Helper: set cell shading ──
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)

def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=None):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Reduce paragraph spacing
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)

# ── Define the RBAC matrix data ──
# Format: (Module/Feature, Admin, Shop Owner, Tourist, Public)
# Symbols: ✓ = full, ◐ = partial/own, ✗ = denied, — = N/A

headers = ['Chức năng / Module', 'Admin', 'Shop Owner', 'Tourist', 'Public']

rows = [
    # Auth
    ('Đăng ký tài khoản',              '—',  '✓',  '✓',  '✓'),
    ('Đăng nhập / Đăng xuất',          '✓',  '✓',  '✓',  '—'),
    ('Đặt lại mật khẩu',              '✓',  '✓',  '✓',  '✓'),
    ('Quản lý hồ sơ cá nhân',         '✓',  '✓',  '✓',  '✗'),
    # POI
    ('Tạo POI',                         '✓',  '◐',  '✗',  '✗'),
    ('Xem danh sách POI (Admin)',       '✓',  '✗',  '✗',  '✗'),
    ('Xem POI sở hữu',                 '✓',  '◐',  '✗',  '✗'),
    ('Sửa POI',                         '✓',  '◐',  '✗',  '✗'),
    ('Xóa POI (soft delete)',           '✓',  '✗',  '✗',  '✗'),
    ('Đổi trạng thái POI',             '✓',  '✗',  '✗',  '✗'),
    ('Upload media cho POI',           '✓',  '◐',  '✗',  '✗'),
    ('Xóa media POI',                  '✓',  '✗',  '✗',  '✗'),
    # Tour
    ('Tạo Tour',                        '✓',  '✗',  '✗',  '✗'),
    ('Sửa Tour / Set POIs',            '✓',  '✗',  '✗',  '✗'),
    ('Xóa Tour',                        '✓',  '✗',  '✗',  '✗'),
    # TTS & QR
    ('Tạo TTS audio',                  '✓',  '◐',  '✗',  '✗'),
    ('Xem giọng đọc TTS',              '✓',  '✓',  '✗',  '✗'),
    ('Xem / Download QR Code',         '✓',  '✓',  '✗',  '✗'),
    ('Sinh lại QR Code',               '✓',  '✗',  '✗',  '✗'),
    # Merchant
    ('Quản lý Merchant (CRUD)',         '✓',  '✗',  '✗',  '✗'),
    # Shop Owner
    ('Xem/Sửa hồ sơ cửa hàng',       '✗',  '✓',  '✗',  '✗'),
    ('Xem analytics riêng',            '✗',  '✓',  '✗',  '✗'),
    # Tourist
    ('Quản lý yêu thích',              '✗',  '✗',  '✓',  '✗'),
    ('Xem lịch sử trải nghiệm',       '✗',  '✗',  '✓',  '✗'),
    # Analytics
    ('Xem analytics tổng quan',         '✓',  '✗',  '✗',  '✗'),
    # Public
    ('Xem POI / Tour (active)',         '—',  '—',  '✓',  '✓'),
    ('Tìm POI gần (nearby)',           '—',  '—',  '✓',  '✓'),
    ('Quét QR Code (validate)',         '—',  '—',  '✓',  '✓'),
    ('Ghi nhật ký trigger',            '—',  '—',  '✓',  '✓'),
]

# ── Group headers (to be inserted as merged-ish bold rows) ──
groups = {
    0: 'Xác thực & Hồ sơ',
    4: 'Quản lý POI',
    12: 'Quản lý Tour',
    15: 'TTS & QR Code',
    19: 'Quản lý Merchant',
    20: 'Shop Owner Portal',
    22: 'Tourist Portal',
    24: 'Analytics',
    25: 'Public API (không cần auth)',
}

# ── Build the rows with group headers ──
final_rows = []
for i, row in enumerate(rows):
    if i in groups:
        final_rows.append(('__GROUP__', groups[i]))
    final_rows.append(row)

# ── Create the table ──
num_cols = 5
num_rows = 1 + len(final_rows)  # header + data rows
table = doc.add_table(rows=num_rows, cols=num_cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set table borders manually (no built-in Table Grid style in this doc)
def set_table_borders(tbl):
    tblPr = tbl._element.tblPr if tbl._element.tblPr is not None else tbl._element.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): 'B0B0B0',
        })
        borders.append(el)
    tblPr.append(borders)
    if tbl._element.tblPr is None:
        tbl._element.insert(0, tblPr)

set_table_borders(table)

# Set column widths
col_widths = [Cm(6.5), Cm(2.0), Cm(2.5), Cm(2.0), Cm(2.0)]
for i, width in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = width

# ── Fill header row ──
header_row = table.rows[0]
for j, header_text in enumerate(headers):
    cell = header_row.cells[j]
    set_cell_text(cell, header_text, bold=True, size=9, color=(255, 255, 255))
    set_cell_shading(cell, '0C4A6E')  # Ocean blue

# ── Fill data rows ──
row_idx = 1
for item in final_rows:
    if item[0] == '__GROUP__':
        # Group header row - merge all cells and style
        row = table.rows[row_idx]
        # Merge cells
        merged = row.cells[0].merge(row.cells[4])
        set_cell_text(merged, item[1], bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=9, color=(12, 74, 110))
        set_cell_shading(merged, 'E0F2FE')  # Light blue
    else:
        row = table.rows[row_idx]
        # Feature name
        set_cell_text(row.cells[0], item[0], bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)

        # Permission cells
        for j in range(1, 5):
            cell = row.cells[j]
            val = item[j]
            if val == '✓':
                set_cell_text(cell, '✓', bold=True, size=10, color=(22, 163, 74))  # green
            elif val == '◐':
                set_cell_text(cell, '◐', bold=True, size=10, color=(234, 138, 0))  # orange
            elif val == '✗':
                set_cell_text(cell, '✗', bold=True, size=10, color=(220, 38, 38))  # red
            elif val == '—':
                set_cell_text(cell, '—', bold=False, size=9, color=(156, 163, 175))  # gray

        # Alternate row shading
        if row_idx % 2 == 0:
            for j in range(5):
                if not (item[0] == '__GROUP__'):
                    set_cell_shading(row.cells[j], 'F8FAFC')  # very light gray

    row_idx += 1

# ── Add legend paragraph after the table ──
# We need to move the table to after the correct paragraph
# First, remove the table from its current position (end of document)
table_element = table._element
table_element.getparent().remove(table_element)

# Insert a description paragraph before the table
ref_element = doc.paragraphs[insert_after_idx]._element

# Add intro paragraph
from docx.oxml.ns import qn as _qn
from copy import deepcopy

def make_paragraph(text, bold=False, size=9):
    """Create a new paragraph element."""
    new_p = doc.paragraphs[0]._element.makeelement(_qn('w:p'), {})
    pPr = new_p.makeelement(_qn('w:pPr'), {})
    pStyle = new_p.makeelement(_qn('w:pStyle'), {_qn('w:val'): 'Normal'})
    pPr.append(pStyle)
    new_p.insert(0, pPr)
    r = new_p.makeelement(_qn('w:r'), {})
    rPr = new_p.makeelement(_qn('w:rPr'), {})
    if bold:
        b = new_p.makeelement(_qn('w:b'), {})
        rPr.append(b)
    sz = new_p.makeelement(_qn('w:sz'), {_qn('w:val'): str(size * 2)})
    rPr.append(sz)
    r.append(rPr)
    t = new_p.makeelement(_qn('w:t'), {})
    t.text = text
    t.set(_qn('xml:space'), 'preserve')
    r.append(t)
    new_p.append(r)
    return new_p

# Insert in reverse order (each addnext goes right after ref)
legend_text = 'Chú thích: ✓ = Toàn quyền    ◐ = Chỉ dữ liệu sở hữu (owner_id = user_id)    ✗ = Không có quyền (HTTP 403)    — = Không áp dụng'
legend_p = make_paragraph(legend_text, size=8)

intro_text = 'Bảng dưới đây mô tả chi tiết quyền truy cập của từng vai trò trên 54 API endpoints, tổ chức theo 9 nhóm chức năng:'
intro_p = make_paragraph(intro_text, size=9)

# Insert: ref → intro_p → table → legend (insert in reverse)
ref_element.addnext(legend_p)
ref_element.addnext(table_element)
ref_element.addnext(intro_p)

print('  + Inserted RBAC matrix table with 9 groups, 29 feature rows')
print('  + Inserted legend paragraph')

# ── Save ──
doc.save(DST)
print(f'\nSaved to {DST}')
