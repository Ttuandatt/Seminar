"""
Fill content into 'Báo cáo Seminar.docx' skeleton.
Run: python -X utf8 docs/fill_report.py
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import os, re

SRC = os.path.join(os.path.dirname(__file__), 'Báo cáo Seminar.docx')
DST = os.path.join(os.path.dirname(__file__), 'Báo cáo Seminar - filled.docx')

doc = Document(SRC)

# ── helpers ──────────────────────────────────────────────
def find_para_index(text_fragment):
    """Return index of first paragraph whose text contains fragment."""
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i
    return None

def insert_after(index, texts, style='Normal'):
    """Insert list of text paragraphs after doc.paragraphs[index]."""
    ref = doc.paragraphs[index]._element
    for txt in reversed(texts):          # reverse because each addnext goes right after ref
        new_p = deepcopy(doc.paragraphs[0]._element)  # clone a blank <w:p>
        new_p.clear()
        from docx.oxml.ns import qn
        r = new_p.makeelement(qn('w:r'), {})
        # Apply style via pPr
        pPr = new_p.makeelement(qn('w:pPr'), {})
        pStyle = new_p.makeelement(qn('w:pStyle'), {qn('w:val'): style})
        pPr.append(pStyle)
        new_p.insert(0, pPr)
        # Handle multi-line: split by \n
        lines = txt.split('\n')
        for li, line in enumerate(lines):
            r2 = new_p.makeelement(qn('w:r'), {})
            t = new_p.makeelement(qn('w:t'), {})
            t.text = line
            t.set(qn('xml:space'), 'preserve')
            r2.append(t)
            new_p.append(r2)
            if li < len(lines) - 1:
                br = new_p.makeelement(qn('w:r'), {})
                br_el = new_p.makeelement(qn('w:br'), {})
                br.append(br_el)
                new_p.append(br)
        ref.addnext(new_p)

def add_content(heading_text, paragraphs, style='Normal'):
    """Find heading, insert content paragraphs after it."""
    idx = find_para_index(heading_text)
    if idx is None:
        print(f'  WARNING: heading not found: {heading_text}')
        return
    insert_after(idx, paragraphs, style)
    print(f'  + {heading_text[:60]}')

# ═══════════════════════════════════════════════════════════
# A. TỔNG QUAN VỀ ỨNG DỤNG VÀ ĐẶC TẢ YÊU CẦU
# ═══════════════════════════════════════════════════════════

print('Section A...')

add_content('1. Lý do chọn đề tài', [
    'Khu phố ẩm thực Vĩnh Khánh (Quận 4, TP. HCM) là điểm đến nổi tiếng với hơn 100 quán ăn trải dọc hai bên đường, thu hút lượng lớn du khách trong và ngoài nước. Tuy nhiên, du khách hiện tại thiếu thông tin chi tiết về từng địa điểm, đặc biệt là du khách quốc tế gặp rào cản ngôn ngữ. Việc quản lý nội dung đa địa điểm vẫn đang dùng phương pháp thủ công (Excel, Word), gây mất thời gian và dễ sai sót.',
    'Bên cạnh đó, các đơn vị quản lý du lịch tại khu vực thiếu dữ liệu số hóa về hành vi của du khách (điểm nào được ghé nhiều, thời gian lưu trú, ngôn ngữ ưa thích) để cải thiện dịch vụ. Do đó, nhóm lựa chọn đề tài xây dựng hệ thống thuyết minh du lịch tự động GPS Tours nhằm giải quyết đồng thời các vấn đề trên thông qua công nghệ định vị GPS và audio guide đa ngôn ngữ.',
])

add_content('2. Mục tiêu', [
    'Đồ án hướng đến các mục tiêu chính sau:',
    '- Xây dựng ứng dụng di động (Tourist App) cho phép du khách nhận thuyết minh âm thanh tự động khi đến gần các điểm tham quan (POI) dựa trên GPS.',
    '- Xây dựng Admin Dashboard trên web để quản trị viên quản lý toàn bộ POIs, Tours, Users và nội dung đa ngôn ngữ (Việt/Anh).',
    '- Xây dựng Shop Owner Dashboard để chủ quán tự quản lý thông tin cửa hàng và POI của mình.',
    '- Thiết kế RESTful API backend với xác thực JWT, phân quyền theo vai trò (RBAC), và tích hợp Text-to-Speech tự động.',
    '- Hỗ trợ chế độ offline thông qua QR Code fallback khi GPS không chính xác.',
    '- Thu thập dữ liệu analytics (lượt xem, lượt nghe audio, trigger GPS/QR) để phục vụ phân tích hành vi du khách.',
])

add_content('3. Phạm vi', [
    'Phạm vi MVP của đồ án bao gồm ba nền tảng chính:',
    '- Admin Dashboard (Web): Đăng nhập, quản lý POI (CRUD), quản lý Tour, quản lý Shop Owner, upload media, tạo TTS tự động, xem bản đồ tổng quan, và xem Analytics.',
    '- Shop Owner Dashboard (Web): Đăng ký/đăng nhập, quản lý POI sở hữu (thêm/sửa), upload hình ảnh và audio, tạo TTS, xem thống kê riêng.',
    '- Tourist App (Mobile): Kiểm tra cấu hình thiết bị (GPS + Internet), bản đồ tương tác, auto-trigger audio theo GPS, theo dõi Tour, quét QR Code, yêu thích, lịch sử, đa ngôn ngữ (VI/EN), chế độ offline.',
    'Khu vực triển khai: Phố ẩm thực Vĩnh Khánh, Quận 4, TP. Hồ Chí Minh.',
    'Ngoài phạm vi: Tính năng AR, Chatbot AI, Social features, thanh toán trực tuyến (dự kiến Phase 3).',
])

add_content('4. Phương pháp phát triển', [
    'Nhóm áp dụng quy trình phát triển Agile-Scrum với các đặc điểm:',
    '- Sprint: 1-2 tuần/sprint, bắt đầu từ Phase 0 (Requirements) đến Phase 6 (Deployment).',
    '- Công cụ quản lý: Git + GitHub cho version control, Claude Code làm AI pair-programming assistant.',
    '- Tài liệu: 13 PRD documents được viết trước khi code, bao gồm Executive Summary, Scope Definition, User Personas, User Stories, Functional Requirements, Acceptance Criteria, NFRs, Data Requirements, API Specifications, UI/UX Specs, Business Rules, Technical Constraints, và Dependencies & Risks.',
    '- Review: Code review qua Pull Request, smoke test end-to-end sau mỗi sprint.',
    '- Tổng thời gian MVP ước tính: 8-10 tuần.',
])

# ── II. TỔNG QUAN VỀ PHẦN MỀM ──

add_content('1.1. Hiệu năng', [
    'Hệ thống đặt ra các chỉ tiêu hiệu năng sau:',
    '- API response time (p95): dưới 500ms cho tất cả các endpoint.',
    '- API response time (p99): dưới 1 giây.',
    '- First Contentful Paint (FCP) trên Admin Dashboard: dưới 2 giây.',
    '- Thời gian render bản đồ trên mobile: dưới 1 giây.',
    '- Độ trễ bắt đầu phát audio (audio start latency): dưới 500ms.',
    '- Hỗ trợ đồng thời tối thiểu 1.000 người dùng cùng lúc.',
    '- Lazy loading hình ảnh để tối ưu băng thông.',
])

add_content('1.2. Bảo mật', [
    'Các yêu cầu bảo mật chính:',
    '- Toàn bộ traffic sử dụng HTTPS.',
    '- Mật khẩu được hash bằng bcrypt với cost factor 12.',
    '- JWT token: Access token hết hạn sau 15 phút, Refresh token sau 7 ngày.',
    '- Rate limiting: 100 req/min/IP (public), 200/min (Shop Owner), 300/min (Admin).',
    '- Validation input bắt buộc phía server (Prisma parameterized queries chống SQL Injection).',
    '- Sanitize output chống XSS, cấu hình CORS whitelist.',
    '- Cô lập dữ liệu Shop Owner: mọi query đều filter theo owner_id, trả về 403 nếu truy cập dữ liệu người khác.',
    '- Phân quyền RBAC: mỗi endpoint được bảo vệ bởi JwtAuthGuard + RolesGuard.',
])

add_content('1.3. Khả năng tương thích', [
    '- Admin Dashboard: Hỗ trợ Chrome 90+, Firefox 90+, Safari 14+, Edge 90+.',
    '- Tourist App: iOS 14+ và Android 10+ thông qua React Native (Expo).',
    '- Responsive design: 320px đến 2560px.',
    '- Backward compatibility: API versioning (/api/v1), deprecation notice 3 tháng.',
])

add_content('1.4. Môi trường', [
    '- Development: Node.js v24, Docker Desktop (PostgreSQL 15 + Redis 7), VS Code.',
    '- Mobile testing: Expo Go trên thiết bị thực tế (Android/iOS).',
    '- Database management: Prisma Studio, Prisma Migrate.',
    '- API testing: Swagger UI tích hợp sẵn tại /api.',
    '- Deployment target: Render.com (backend), EAS Build (mobile APK/IPA).',
    '- Timezone mặc định: Asia/Ho_Chi_Minh, định dạng ngày: DD/MM/YYYY.',
])

add_content('1.5. Khả năng mở rộng và bảo trì', [
    '- Kiến trúc module-based (NestJS modules): mỗi domain nghiệp vụ là một module độc lập, dễ thêm/bớt.',
    '- TypeScript strict mode toàn bộ codebase đảm bảo type safety.',
    '- ESLint + Prettier enforced cho code formatting thống nhất.',
    '- Mục tiêu test coverage > 70%.',
    '- CI/CD pipeline hỗ trợ automated build và deploy, rollback dưới 5 phút.',
    '- Hỗ trợ horizontal scaling khi cần mở rộng.',
])

add_content('2.1. Admin Dashboard', [
    'Admin Dashboard là ứng dụng web React cung cấp các tính năng quản trị:',
    '- Đăng nhập/đăng xuất với JWT authentication.',
    '- Dashboard tổng quan: hiển thị số liệu POIs, Tours, Users, lượt xem.',
    '- Quản lý POI: danh sách có phân trang, tìm kiếm, lọc theo category/status; form tạo/sửa POI với map picker chọn vị trí, upload hình ảnh và audio, tạo TTS tự động từ nội dung mô tả.',
    '- Quản lý Tour: tạo/sửa tour, thiết lập danh sách POI trong tour với thứ tự tùy chỉnh.',
    '- Quản lý Merchant (Shop Owner): tạo tài khoản, xem danh sách, chỉnh sửa.',
    '- Bản đồ tổng quan: hiển thị tất cả POI trên Google Maps, phân biệt theo category.',
    '- Analytics: thống kê tổng quan, top POIs theo lượt xem, phân bố trigger (GPS/QR/Manual).',
    '- Quản lý hồ sơ cá nhân.',
])

add_content('2.2. Tourist App', [
    'Tourist App là ứng dụng di động React Native (Expo) dành cho du khách:',
    '- Kiểm tra cấu hình thiết bị: verify GPS permission và Internet connectivity khi khởi động.',
    '- Bản đồ tương tác: hiển thị các POI active trên bản đồ, định vị người dùng real-time.',
    '- Auto-trigger GPS: khi du khách đi vào bán kính trigger của POI (mặc định 15m), tự động hiển thị thông tin và phát audio thuyết minh.',
    '- Danh sách Tour: xem và theo dõi tour có hướng dẫn (turn-by-turn navigation giữa các POI).',
    '- QR Scanner: quét mã QR tại địa điểm để nghe thuyết minh (fallback khi GPS không chính xác).',
    '- Yêu thích và Lịch sử: lưu POI yêu thích, xem lịch sử trải nghiệm.',
    '- Đa ngôn ngữ: chuyển đổi Tiếng Việt/English, nội dung POI và audio theo ngôn ngữ đã chọn.',
    '- Đăng ký/Đăng nhập (tùy chọn): để đồng bộ yêu thích và lịch sử.',
    '- Chế độ Offline: đồng bộ dữ liệu POI về máy, quét QR offline.',
])

add_content('2.3. Shop Owner Dashboard', [
    'Shop Owner Dashboard là giao diện web dành cho chủ quán/cửa hàng:',
    '- Đăng ký tài khoản Shop Owner (email, mật khẩu, thông tin cửa hàng).',
    '- Dashboard: xem tổng quan POI sở hữu, lượt xem, lượt nghe audio.',
    '- Quản lý POI: tạo mới, chỉnh sửa POI (tên, mô tả VI/EN, vị trí); upload hình ảnh và audio.',
    '- Tạo TTS tự động: sinh audio thuyết minh từ nội dung mô tả cho cả Tiếng Việt và Tiếng Anh.',
    '- Analytics: thống kê lượt xem và audio plays chỉ cho POI của mình.',
    '- Quản lý hồ sơ: cập nhật tên quán, địa chỉ, số điện thoại.',
    '- Cô lập dữ liệu: chỉ thấy và thao tác được trên POI có ownerId trùng userId của mình.',
])

add_content('3.1. Các persona người dùng', [
    'Hệ thống phục vụ 3 nhóm người dùng chính:',
    '- Admin (Quản trị viên): Quản lý toàn bộ nội dung, tài khoản và cấu hình hệ thống. Có toàn quyền CRUD trên POIs, Tours, Users. Xem analytics toàn hệ thống.',
    '- Shop Owner (Chủ quán): Quản lý thông tin cửa hàng và POI sở hữu. Có thể tạo/sửa POI, upload media, tạo TTS. Chỉ thấy dữ liệu của mình.',
    '- Tourist (Du khách): Sử dụng app di động để khám phá. Xem bản đồ, nghe audio tự động, theo dõi tour. Đăng nhập tùy chọn để lưu yêu thích và lịch sử.',
])

add_content('3.2. Ma trận phân quyền', [
    'Ma trận phân quyền theo vai trò (RBAC):',
    '- Admin: Toàn quyền trên tất cả modules (POI CRUD, Tour CRUD, User management, Analytics, Merchant management).',
    '- Shop Owner: Xem/Sửa POI sở hữu (owner_id = user_id), Upload media cho POI mình, Tạo TTS, Xem analytics riêng, Cập nhật hồ sơ cửa hàng. KHÔNG được xóa POI, KHÔNG truy cập POI người khác (HTTP 403).',
    '- Tourist: Xem tất cả POI/Tour active (qua Public API không cần auth), Đăng ký/Đăng nhập, Quản lý yêu thích và lịch sử (cần auth), Trigger GPS/QR events. KHÔNG tạo/sửa nội dung.',
])

# ── III. QUY TRÌNH NGHIỆP VỤ ──

add_content('1. Quy trình quản lý nội dung (Admin)', [
    'Quy trình quản lý nội dung POI của Admin gồm các bước:',
    '1. Admin đăng nhập vào Dashboard bằng email/password → nhận JWT token.',
    '2. Tạo POI mới: nhập tên (VI/EN), mô tả, chọn vị trí trên bản đồ (map picker), chọn category, thiết lập trigger radius.',
    '3. Upload media: thêm hình ảnh (tối đa 10 ảnh, mỗi ảnh ≤ 5MB) và audio thuyết minh.',
    '4. Tạo TTS tự động: hệ thống sinh audio từ nội dung mô tả bằng Microsoft Edge TTS (giọng vi-VN-HoaiMyNeural cho Tiếng Việt, en-US-AriaNeural cho Tiếng Anh).',
    '5. Tạo QR Code: hệ thống tự động sinh mã QR (format: gpstours:poi:<poiId>) với error correction level H.',
    '6. Publish: chuyển trạng thái POI từ DRAFT sang ACTIVE để hiển thị trên app.',
    '7. Tạo Tour: tạo tour mới, thêm các POI vào tour theo thứ tự mong muốn.',
    '8. Monitoring: xem analytics tổng quan (lượt xem, audio plays, trigger stats).',
])

add_content('2. Quy trình đăng ký và quản lý cửa hàng (Shop Owner)', [
    '1. Shop Owner truy cập trang đăng ký, nhập email, mật khẩu, họ tên, tên quán, địa chỉ, số điện thoại.',
    '2. Hệ thống tạo tài khoản với role SHOP_OWNER và bản ghi ShopOwner profile liên kết.',
    '3. Sau đăng nhập, Shop Owner thấy Dashboard riêng với danh sách POI sở hữu.',
    '4. Tạo POI mới: nhập thông tin, chọn vị trí, hệ thống tự gán ownerId = userId của Shop Owner.',
    '5. Upload media và tạo TTS cho POI của mình.',
    '6. Xem analytics: số lượt xem và audio plays chỉ cho POI mình sở hữu.',
    '7. Cập nhật hồ sơ cửa hàng: tên quán, địa chỉ, số điện thoại.',
    'Lưu ý: Mọi truy vấn đều filter theo owner_id, nếu Shop Owner cố truy cập POI người khác sẽ nhận HTTP 403 Forbidden.',
])

add_content('3. Quy trình tham quan của du khách (Tourist)', [
    '1. Du khách mở ứng dụng → kiểm tra cấu hình thiết bị (GPS + Internet). Nếu chưa đủ điều kiện, hiển thị màn hình hướng dẫn bật GPS/kết nối mạng.',
    '2. Màn hình chính hiển thị bản đồ với tất cả POI active, vị trí hiện tại được cập nhật real-time.',
    '3. Khi đi vào vùng trigger của POI (bán kính mặc định 15m), app tự động hiển thị thông tin POI gần nhất.',
    '4. Audio thuyết minh được phát tự động (nếu bật auto-play) theo ngôn ngữ đã chọn.',
    '5. Du khách có thể xem chi tiết POI, nghe lại audio, xem hình ảnh.',
    '6. Đăng nhập (tùy chọn) để lưu POI yêu thích và xem lịch sử trải nghiệm.',
    '7. Theo dõi Tour: chọn tour, app hướng dẫn turn-by-turn đến từng POI theo thứ tự.',
])

add_content('4. Quy trình fallback khi không có GPS', [
    'Khi GPS không khả dụng hoặc không chính xác (trong nhà, khu vực đông đúc):',
    '1. Du khách mở chức năng QR Scanner từ menu "Thêm".',
    '2. Quét mã QR được đặt tại mỗi điểm tham quan (mã QR chứa: gpstours:poi:<poiId>).',
    '3. App gửi mã đến server để validate (POST /public/qr/validate).',
    '4. Server kiểm tra format hợp lệ và POI tồn tại với status ACTIVE.',
    '5. Nếu hợp lệ → hiển thị chi tiết POI và phát audio thuyết minh.',
    '6. Nếu đang offline → app kiểm tra dữ liệu đã đồng bộ trong SQLite local, hiển thị từ cache.',
    '7. TriggerLog được ghi nhận với triggerType = QR.',
])

add_content('5. Quy trình theo dõi Tour có hướng dẫn', [
    '1. Du khách chọn một Tour từ danh sách → xem thông tin tour và danh sách điểm dừng.',
    '2. Bấm "Bắt đầu Tour" → app chuyển sang chế độ theo dõi (Tour Follow mode).',
    '3. App hiển thị điểm đến tiếp theo, khoảng cách và hướng di chuyển.',
    '4. Khi đến gần POI (trong bán kính trigger), app đánh dấu "Đã tới nơi!" và phát audio.',
    '5. Du khách bấm "Đến trạm tiếp theo" để chuyển sang POI kế tiếp trong tour.',
    '6. Khi hoàn thành tất cả các điểm dừng → hiển thị màn hình "Đã hoàn thành Tour!" với thông báo chúc mừng.',
    '7. Du khách có thể kết thúc tour sớm bằng nút "Kết thúc Tour".',
])

add_content('6. Quy trình xử lý vùng giao thoa POI', [
    'Khi du khách đứng ở vị trí mà nhiều POI có vùng trigger chồng lấn:',
    '1. App tính khoảng cách từ vị trí hiện tại đến tất cả POI bằng công thức Haversine.',
    '2. Lọc các POI nằm trong bán kính trigger (distanceM ≤ triggerRadius).',
    '3. Sắp xếp theo khoảng cách tăng dần (POI gần nhất ưu tiên cao nhất).',
    '4. Xây dựng nearbyQueue: danh sách POI gần, hiển thị POI gần nhất trước.',
    '5. Audio chỉ phát cho POI ở đầu queue (Audio Singleton: chỉ 1 audio phát tại 1 thời điểm).',
    '6. Khi du khách di chuyển, queue được cập nhật real-time (mỗi 2 giây hoặc khi di chuyển 5m).',
])

# ── IV. PHÂN TÍCH VÀ THIẾT KẾ ──

add_content('1.1. Sơ đồ Use Case tổng quát', [
    'Hệ thống có 3 actor chính: Admin, Shop Owner, Tourist. Sơ đồ Use Case tổng quát mô tả các nhóm chức năng:',
    '- Admin: UC-01 Đăng nhập, UC-11 Tạo POI, UC-12 Sửa POI, UC-13 Xóa POI, UC-14 Quản lý trạng thái POI, UC-21 Tạo Tour, UC-22 Sửa Tour, UC-23 Xóa Tour, UC-31 Xem Analytics, UC-32 Quản lý Merchant.',
    '- Shop Owner: UC-40 Đăng ký/Đăng nhập, UC-41 Quản lý POI sở hữu, UC-42 Upload media, UC-43 Tạo TTS, UC-44 Xem Analytics riêng.',
    '- Tourist: UC-30 Xem bản đồ, UC-51 Auto-trigger nội dung, UC-52 Xử lý vùng giao thoa, UC-53 Quét QR Code, UC-54 Theo dõi Tour, UC-55 Quản lý yêu thích/lịch sử.',
    '(Sơ đồ Use Case chi tiết xem tại docs/step3_prd/14_usecase_diagram.md)',
])

add_content('1.2. UC-01: Đăng nhập Admin', [
    'Actor: Admin',
    'Mô tả: Admin đăng nhập vào hệ thống Dashboard bằng email và password.',
    'Tiền điều kiện: Admin đã có tài khoản trong hệ thống.',
    'Luồng chính: (1) Admin mở trang đăng nhập → (2) Nhập email và password → (3) Hệ thống kiểm tra thông tin → (4) Nếu hợp lệ, tạo JWT access token (15 phút) và refresh token (7 ngày) → (5) Chuyển hướng đến Dashboard.',
    'Luồng ngoại lệ: (3a) Sai thông tin → hiển thị lỗi "Email hoặc mật khẩu không đúng". (3b) Tài khoản bị khóa (5 lần sai liên tiếp) → hiển thị thông báo khóa 30 phút.',
    'Hậu điều kiện: JWT token được lưu, failedLoginCount reset về 0.',
])

add_content('1.3. UC-02: Đăng ký tài khoản', [
    'Actor: Shop Owner, Tourist',
    'Mô tả: Người dùng mới đăng ký tài khoản để sử dụng hệ thống.',
    'Luồng chính: (1) Người dùng chọn "Đăng ký" → (2) Nhập họ tên, email, mật khẩu → (3) Nếu là Shop Owner: nhập thêm tên quán, địa chỉ, SĐT → (4) Hệ thống validate dữ liệu → (5) Tạo tài khoản với role tương ứng → (6) Chuyển về trang đăng nhập.',
    'Luồng ngoại lệ: (4a) Email đã tồn tại → báo lỗi. (4b) Mật khẩu yếu (< 8 ký tự, thiếu chữ hoa/thường/số) → báo lỗi.',
    'Business Rules: BR-102 (Password hash bcrypt), password phải ≥ 8 ký tự, chứa ít nhất 1 chữ hoa, 1 chữ thường, 1 số.',
])

add_content('1.4. UC-11: Tạo POI mới', [
    'Actor: Admin',
    'Mô tả: Admin tạo mới một Point of Interest trong hệ thống.',
    'Luồng chính: (1) Admin chọn "Tạo POI" → (2) Nhập thông tin: tên VI/EN, mô tả VI/EN, chọn category → (3) Chọn vị trí trên map picker hoặc nhập tọa độ → (4) Thiết lập trigger radius (mặc định 15m) → (5) Upload hình ảnh và audio (tùy chọn) → (6) Bấm Lưu → (7) Hệ thống tạo POI (status: DRAFT), tự động sinh TTS từ mô tả, tự động sinh QR Code.',
    'Hậu điều kiện: POI được tạo với status DRAFT. TTS audio được sinh fire-and-forget. QR Code được sinh tự động.',
    'Business Rules: BR-201 (Tên unique), BR-202 (Bắt buộc name_vi + description_vi), BR-203 (Tọa độ trong phạm vi).',
])

add_content('1.5. UC-13: Xóa POI', [
    'Actor: Admin, Shop Owner (chỉ POI sở hữu)',
    'Mô tả: Xóa mềm (soft delete) một POI khỏi hệ thống.',
    'Luồng chính: (1) Chọn POI từ danh sách → (2) Bấm "Xóa" → (3) Xác nhận → (4) Hệ thống kiểm tra quyền → (5) Cập nhật deletedAt = now() (soft delete).',
    'Luồng ngoại lệ: (4a) Shop Owner cố xóa POI không sở hữu → HTTP 403 Forbidden.',
    'Hậu điều kiện: POI bị ẩn khỏi tất cả danh sách (query filter deletedAt = null). Dữ liệu vẫn tồn tại trong DB.',
])

add_content('1.6. UC-21: Tạo Tour', [
    'Actor: Admin',
    'Mô tả: Tạo tour tham quan gồm danh sách POI theo thứ tự.',
    'Luồng chính: (1) Admin chọn "Tạo Tour" → (2) Nhập tên VI/EN, mô tả, thời lượng ước tính → (3) Chọn các POI vào tour, sắp xếp thứ tự → (4) Lưu → (5) Hệ thống tạo Tour (status: DRAFT) và các bản ghi TourPoi liên kết.',
    'Hậu điều kiện: Tour được tạo với danh sách POI có orderIndex.',
])

add_content('1.7. UC-30: Xem bản đồ', [
    'Actor: Tourist',
    'Mô tả: Du khách xem bản đồ tương tác hiển thị các điểm tham quan.',
    'Luồng chính: (1) Mở app → tab Bản đồ → (2) App yêu cầu quyền GPS → (3) Hiển thị bản đồ với vị trí hiện tại và tất cả POI active → (4) POI hiển thị dưới dạng marker, nhấn để xem thông tin tóm tắt → (5) Bấm "Xem chi tiết" để vào trang chi tiết POI.',
    'Business Rules: Chỉ hiển thị POI có status = ACTIVE và deletedAt = null.',
])

add_content('1.8. UC-40: Quản lý POI (Shop Owner)', [
    'Actor: Shop Owner',
    'Mô tả: Shop Owner quản lý các POI sở hữu (có ownerId = userId).',
    'Luồng chính: (1) Đăng nhập → (2) Xem danh sách POI sở hữu → (3) Tạo mới hoặc sửa POI → (4) Upload media (hình ảnh, audio) → (5) Tạo TTS tự động.',
    'Ràng buộc: Mọi query đều filter ownerId = currentUser.id. Truy cập POI người khác → 403 Forbidden.',
])

add_content('1.9. UC-51: Auto-trigger nội dung', [
    'Actor: Tourist',
    'Mô tả: App tự động phát nội dung khi du khách đi vào vùng trigger của POI.',
    'Luồng chính: (1) App theo dõi vị trí GPS real-time (watchPositionAsync, accuracy: High, mỗi 2 giây hoặc 5m) → (2) Tính khoảng cách Haversine đến mỗi POI → (3) Nếu khoảng cách ≤ triggerRadius → thêm vào nearbyQueue → (4) Hiển thị POI gần nhất, phát audio nếu auto-play bật.',
    'Hậu điều kiện: TriggerLog được ghi (triggerType: GPS, userAction, tọa độ, khoảng cách).',
])

add_content('1.10. UC-52: Xử lý vùng giao thoa', [
    'Actor: Hệ thống (tự động)',
    'Mô tả: Khi nhiều POI có vùng trigger chồng lấn, hệ thống chọn POI ưu tiên.',
    'Thuật toán: (1) Tính khoảng cách đến tất cả POI active → (2) Lọc POI trong bán kính trigger → (3) Sắp xếp theo khoảng cách tăng dần → (4) POI gần nhất được ưu tiên hiển thị và phát audio → (5) Queue cập nhật real-time khi du khách di chuyển.',
])

# ── Sơ đồ hoạt động ──

add_content('2.1. AD-01: Luồng đăng nhập Admin', [
    'Activity Diagram mô tả luồng đăng nhập Admin:',
    '(Start) → [Mở trang Login] → [Nhập email + password] → <Validate input> → {Hợp lệ?}',
    '→ [Không] → [Hiển thị lỗi validation] → quay lại nhập',
    '→ [Có] → [Gửi POST /auth/login] → <Server kiểm tra> → {Đúng thông tin?}',
    '→ [Không] → [Tăng failedLoginCount] → {≥ 5 lần?} → [Có] → [Khóa tài khoản 30 phút] → (End)',
    '→ [Không] → [Hiển thị lỗi "Sai email hoặc mật khẩu"] → quay lại nhập',
    '→ [Đúng] → [Tạo JWT access + refresh token] → [Reset failedLoginCount = 0] → [Redirect Dashboard] → (End)',
    '(Sơ đồ chi tiết xem tại docs/step3_prd/16_activity_diagrams.md)',
])

add_content('2.2. AD-02: Vòng đời POI (CRUD)', [
    'Activity Diagram mô tả vòng đời POI từ tạo đến xóa:',
    '(Start) → [Admin tạo POI] → POI(DRAFT) → [Upload media + TTS] → [Admin publish] → POI(ACTIVE)',
    '→ {Cần sửa?} → [Có] → [Admin/SO chỉnh sửa] → [Re-generate TTS nếu mô tả thay đổi] → POI(ACTIVE)',
    '→ {Archive?} → [Có] → POI(ARCHIVED) → [Ẩn khỏi app du khách]',
    '→ {Xóa?} → [Có] → [Soft delete: deletedAt = now()] → [Ẩn khỏi mọi danh sách] → (End)',
    'Trạng thái: DRAFT → ACTIVE ↔ ARCHIVED. Soft delete áp dụng ở mọi trạng thái.',
])

add_content('2.3. AD-04: Hành trình du khách', [
    'Activity Diagram mô tả luồng trải nghiệm du khách:',
    '(Start) → [Mở app] → [Kiểm tra thiết bị (GPS + Internet)] → {Đủ điều kiện?}',
    '→ [Không] → [Hiển thị Device Check screen] → [Bật GPS/kết nối mạng] → [Kiểm tra lại]',
    '→ [Có] → [Hiển thị Landing screen] → {Lần đầu?} → [Có] → [Onboarding 3 slide]',
    '→ [Xong/Bỏ qua] → [Bản đồ chính] → {GPS trigger?}',
    '→ [Có] → [Hiển thị POI gần nhất] → [Phát audio] → [Ghi TriggerLog]',
    '→ [Du khách di chuyển] → [Cập nhật nearbyQueue] → lặp lại kiểm tra GPS trigger',
    '→ [Kết thúc tham quan] → (End)',
])

add_content('2.4. AD-05: Phát hiện vị trí và auto-trigger GPS', [
    'Activity Diagram chi tiết luồng GPS auto-trigger:',
    '(Start) → [watchPositionAsync(accuracy: High, interval: 2s, distance: 5m)]',
    '→ [Nhận vị trí mới (lat, lng)] → [ForEach POI active] → [Tính Haversine distance]',
    '→ {distance ≤ triggerRadius?} → [Có] → [Thêm vào inRange list với distanceM]',
    '→ [Sort inRange theo distance ASC] → [Set nearbyQueue = inRange]',
    '→ {nearbyQueue thay đổi?} → [Có] → [Hiển thị POI đầu queue]',
    '→ {autoPlayAudio bật?} → [Có] → [playGlobalAudio(audioUrl, poiId)] → (Audio Singleton phát)',
    '→ [POST /public/trigger-log (GPS, ACCEPTED, coords, distance)]',
    '→ [Chờ vị trí mới] → lặp lại',
])

add_content('2.5. AD-06: Shop Owner đăng ký và quản lý POI', [
    'Activity Diagram luồng Shop Owner:',
    '(Start) → [Mở trang đăng ký Shop Owner] → [Nhập email, password, họ tên, tên quán, địa chỉ, SĐT]',
    '→ [POST /auth/register (role: SHOP_OWNER)] → {Thành công?}',
    '→ [Không] → [Hiển thị lỗi (email trùng, validation)] → quay lại',
    '→ [Có] → [Tạo User + ShopOwner profile] → [Đăng nhập]',
    '→ [Dashboard Shop Owner] → [Xem danh sách POI (GET /shop-owner/pois)]',
    '→ {Tạo POI mới?} → [Có] → [Nhập thông tin, chọn vị trí] → [POST /shop-owner/pois (ownerId auto-set)]',
    '→ [Upload media] → [Tạo TTS] → [POI hiển thị trên app khi ACTIVE]',
])

# ── Sơ đồ luồng dữ liệu ──

add_content('3.1. DFD mức 0 — Tổng quan toàn hệ thống', [
    'DFD mức 0 gồm 1 process trung tâm "Hệ thống GPS Tours" và 3 external entities:',
    '- Admin → Hệ thống: Quản lý POI, Tour, User, Analytics',
    '- Shop Owner → Hệ thống: Quản lý POI sở hữu, Media, TTS',
    '- Tourist → Hệ thống: Xem bản đồ, Trigger GPS/QR, Xem Tour, Yêu thích, Lịch sử',
    'Data stores: D1-Users, D2-POIs, D3-Tours, D4-Media, D5-TriggerLogs, D6-Favorites, D7-ViewHistory',
])

add_content('3.2. DFD mức 1 — Phân rã các tiến trình chính', [
    'DFD mức 1 phân rã thành 6 process chính:',
    'P1. Xác thực (Authentication): Xử lý đăng nhập, đăng ký, refresh token, reset password. Kết nối D1-Users.',
    'P2. Quản lý nội dung (Content Management): CRUD POI, Tour. Kết nối D2-POIs, D3-Tours, D4-Media.',
    'P3. Xử lý định vị (Location Processing): Nhận tọa độ GPS, tính khoảng cách, xác định POI gần nhất. Kết nối D2-POIs, D5-TriggerLogs.',
    'P4. Phát audio (Audio Delivery): Chọn audio theo ngôn ngữ, phát qua Audio Singleton. Kết nối D4-Media.',
    'P5. Quản lý du khách (Tourist Management): Yêu thích, lịch sử, profile. Kết nối D6-Favorites, D7-ViewHistory.',
    'P6. Phân tích dữ liệu (Analytics): Tổng hợp lượt xem, audio plays, trigger stats. Kết nối D5, D7.',
])

add_content('3.3. DFD mức 2 — Tiến trình định vị và kích hoạt nội dung', [
    'DFD mức 2 chi tiết cho P3 (Xử lý định vị):',
    'P3.1. Theo dõi vị trí: Nhận GPS coordinates từ Tourist device (watchPositionAsync).',
    'P3.2. Tính khoảng cách: Áp dụng công thức Haversine cho mỗi POI active.',
    'P3.3. Lọc vùng trigger: So sánh distance với triggerRadius, lọc POI trong vùng.',
    'P3.4. Sắp xếp ưu tiên: Sort theo khoảng cách tăng dần (Criteria Engine).',
    'P3.5. Ghi nhật ký: Ghi TriggerLog (deviceId, poiId, triggerType, userAction, coordinates, distance).',
    'Input: GPS coordinates từ device. Output: nearbyQueue (danh sách POI sắp xếp), TriggerLog entry.',
])

# ── Sơ đồ tuần tự ──

add_content('4.1. SD-01: Admin đăng nhập (JWT flow)', [
    'Sequence Diagram mô tả luồng đăng nhập JWT:',
    'Admin → Browser: Nhập email + password',
    'Browser → API [POST /auth/login]: {email, password}',
    'API → AuthService: validate(email, password)',
    'AuthService → PrismaDB: findUnique(email)',
    'PrismaDB → AuthService: User record',
    'AuthService → AuthService: bcrypt.compare(password, passwordHash)',
    'AuthService → AuthService: Check failedLoginCount < 5, lockedUntil',
    'AuthService → JwtService: sign({sub: userId, role, jti: uuid}, 15min)',
    'JwtService → AuthService: accessToken + refreshToken',
    'AuthService → API: {accessToken, refreshToken, user}',
    'API → Browser: 200 OK + tokens',
    'Browser → localStorage: Lưu tokens',
    'Browser → Admin: Redirect /admin/dashboard',
])

add_content('4.2. SD-03: Admin tạo POI', [
    'Sequence Diagram luồng tạo POI:',
    'Admin → Browser: Điền form POI + chọn vị trí trên map',
    'Browser → API [POST /pois]: {nameVi, nameEn, descriptionVi, descriptionEn, lat, lng, category, triggerRadius} + Header: Bearer token',
    'API → JwtAuthGuard: Validate token',
    'API → RolesGuard: Check role = ADMIN',
    'API → PoisService.create(): Tạo POI record',
    'PoisService → PrismaDB: poi.create({...dto, createdById})',
    'PrismaDB → PoisService: POI created (status: DRAFT)',
    'PoisService → TtsService (fire-and-forget): autoGenerateTts(poiId, descriptionVi, descriptionEn)',
    'TtsService → MsEdgeTTS: Stream audio generation (vi-VN-HoaiMyNeural)',
    'TtsService → FileSystem: Save to /uploads/tts/',
    'TtsService → PrismaDB: poiMedia.create({type: AUDIO, language: VI})',
    'PoisService → QrService (fire-and-forget): generateForPoi(poiId)',
    'QrService → QRCode: toFile("gpstours:poi:<id>", PNG, 512px, errorLevel: H)',
    'PoisService → SeedExportService (fire-and-forget): exportSeedData()',
    'API → Browser: 201 Created + POI data',
])

add_content('4.3. SD-06: Tourist xem bản đồ và GPS auto-trigger', [
    'Sequence Diagram luồng GPS auto-trigger:',
    'Tourist → App: Mở tab Bản đồ',
    'App → ExpoLocation: watchPositionAsync(High, 2s, 5m)',
    'ExpoLocation → App: Location callback (lat, lng)',
    'App → App: ForEach POI: haversine(userLat, userLng, poiLat, poiLng)',
    'App → App: Filter distance ≤ triggerRadius → sort by distance',
    'App → App: setNearbyQueue(sortedPois)',
    'App → UI: Hiển thị POI card gần nhất',
    'App → AudioContext: playGlobalAudio(audioUrl, poiId) [nếu autoPlay]',
    'AudioContext → ExpoAudio: Load + play audio file',
    'App → API [POST /public/trigger-log]: {deviceId, poiId, triggerType: GPS, userAction: ACCEPTED, lat, lng, distance}',
    'API → PrismaDB: triggerLog.create()',
])

add_content('4.4. SD-09: Đặt lại mật khẩu', [
    'Sequence Diagram luồng Forgot Password:',
    'User → App/Browser: Bấm "Quên mật khẩu"',
    'App → API [POST /auth/forgot-password]: {email}',
    'API → AuthService: generateResetToken(email)',
    'AuthService → PrismaDB: user.findUnique(email)',
    'AuthService → AuthService: Generate UUID token, expiry 1h',
    'AuthService → PrismaDB: passwordResetToken.create({userId, token, expiresAt})',
    'API → User: 200 OK (luôn trả success để không leak email tồn tại)',
    'User → App: Nhập token + mật khẩu mới',
    'App → API [POST /auth/reset-password]: {token, newPassword}',
    'API → AuthService: Verify token valid + not expired + not used',
    'AuthService → PrismaDB: user.update({passwordHash: bcrypt(newPassword)})',
    'AuthService → PrismaDB: passwordResetToken.update({usedAt: now()})',
    'API → User: 200 OK "Mật khẩu đã được đặt lại"',
])

add_content('4.5. SD-16: Audio Singleton', [
    'Sequence Diagram mô tả Audio Singleton pattern:',
    'Component A → AudioContext: playGlobalAudio(urlA, poiIdA)',
    'AudioContext → AudioContext: currentPoiId = poiIdA, isPlaying = true',
    'AudioContext → ExpoAudio: useAudioPlayer.replace(urlA)',
    'ExpoAudio → Speaker: Phát audio A',
    '[Trong khi audio A đang phát...]',
    'Component B → AudioContext: playGlobalAudio(urlB, poiIdB)',
    'AudioContext → ExpoAudio: player.pause() [dừng audio A]',
    'AudioContext → AudioContext: currentPoiId = poiIdB',
    'AudioContext → ExpoAudio: useAudioPlayer.replace(urlB)',
    'ExpoAudio → Speaker: Phát audio B',
    'Đảm bảo: Chỉ có 1 audio phát tại bất kỳ thời điểm nào (Singleton pattern).',
])

add_content('4.6. SD-17: QR Code Offline Fallback', [
    'Sequence Diagram luồng QR offline:',
    'Tourist → App: Mở QR Scanner',
    'App → Camera: requestPermission() → scan QR code',
    'Camera → App: QR data = "gpstours:poi:<poiId>"',
    'App → App: Parse poiId from QR string',
    '{Online?}',
    '→ [Có] → App → API [POST /public/qr/validate]: {data: "gpstours:poi:<id>"}',
    '  API → PrismaDB: poi.findFirst({id, status: ACTIVE})',
    '  API → App: 200 OK + full POI data with media',
    '  App → SQLite: Cache POI data locally',
    '→ [Không] → App → SQLite: SELECT * FROM pois WHERE id = poiId',
    '  SQLite → App: Cached POI data (nếu đã sync trước đó)',
    'App → UI: Hiển thị chi tiết POI',
    'App → AudioContext: Phát audio (nếu có cached)',
])

print('Section A done.')

# ═══════════════════════════════════════════════════════════
# B. KIẾN TRÚC PHẦN MỀM
# ═══════════════════════════════════════════════════════════

print('Section B...')

add_content('1. Kiến trúc N-tier module-based', [
    'Hệ thống GPS Tours được thiết kế theo kiến trúc N-tier kết hợp module-based, chia thành 4 tầng chính: Client, API Gateway, Service, và Data. Kiến trúc này đảm bảo separation of concerns, dễ bảo trì và mở rộng từng tầng độc lập.',
])

add_content('1.1. Tầng Client', [
    'Tầng Client gồm 3 ứng dụng giao tiếp với backend qua RESTful API:',
    '- Admin Dashboard: Ứng dụng React 19 (Vite) với React Router v7, Tailwind CSS, Leaflet Maps. Chạy trên trình duyệt web. Gửi request kèm JWT Bearer token.',
    '- Shop Owner Dashboard: Tích hợp chung với Admin Dashboard, sử dụng routing riêng (/owner/*) và layout riêng (ShopOwnerLayout). Phân biệt bằng role trong JWT.',
    '- Tourist App: Ứng dụng React Native (Expo SDK 54) chạy trên iOS/Android. Sử dụng Expo Go trong phát triển, EAS Build cho production. Giao tiếp qua API tương tự.',
])

add_content('1.2. Tầng API Gateway', [
    'Tầng API Gateway do NestJS 11 đảm nhiệm, cung cấp:',
    '- Routing: Phân luồng request đến đúng module xử lý (AuthModule, PoisModule, ToursModule, v.v.).',
    '- Authentication: JwtAuthGuard validate JWT token trên mọi request protected.',
    '- Authorization: RolesGuard kiểm tra role (ADMIN, SHOP_OWNER, TOURIST) dựa trên decorator @Roles().',
    '- Validation: DTOs với class-validator đảm bảo input hợp lệ trước khi đến service layer.',
    '- Swagger UI: Tự động sinh API documentation tại /api.',
    '- Static file serving: ServeStaticModule phục vụ thư mục /uploads (hình ảnh, audio, QR code).',
])

add_content('1.3. Tầng Service', [
    'Tầng Service chứa business logic, tổ chức theo NestJS modules:',
    '- AuthModule: Xử lý đăng nhập, đăng ký, refresh token, reset password. Sử dụng Passport.js + JWT Strategy.',
    '- PoisModule: CRUD POI cho Admin, tích hợp TtsModule và QrModule để auto-generate TTS audio và QR code.',
    '- ToursModule: CRUD Tour, quản lý danh sách POI trong tour (TourPoi junction table).',
    '- ShopOwnerModule: Portal cho Shop Owner, filter data theo ownerId, hỗ trợ media upload và TTS.',
    '- TouristModule: Quản lý profile, favorites, view history cho du khách đã đăng nhập.',
    '- PublicModule: API công khai (không cần auth) cho mobile app — lấy POI, Tour, validate QR, log trigger.',
    '- TtsModule: Tích hợp Microsoft Edge TTS (msedge-tts) để sinh audio từ text.',
    '- QrModule: Sinh mã QR Code (PNG, 512px, error correction H) cho mỗi POI.',
    '- MediaModule: Upload và quản lý file (Multer, UUID filename, giới hạn 50MB).',
    '- AnalyticsModule: Tổng hợp thống kê lượt xem, audio plays, trigger distribution.',
    '- SeedExportModule: Tự động export DB ra JSON khi data thay đổi, phục vụ đồng bộ seed data giữa các dev.',
])

add_content('1.4. Tầng Data', [
    'Tầng Data gồm:',
    '- PostgreSQL 15: Cơ sở dữ liệu quan hệ chính, chạy trên Docker Desktop trong development.',
    '- Prisma ORM 5.22: Object-Relational Mapping, cung cấp type-safe queries, migration management, và Prisma Studio.',
    '- File System: Thư mục /uploads lưu trữ hình ảnh, audio, TTS files, QR codes.',
    '- SQLite (Mobile): Cơ sở dữ liệu local trên thiết bị di động để hỗ trợ chế độ offline.',
])

add_content('1. NestJS vs Express', [
    'Nhóm chọn NestJS thay vì Express thuần vì các lý do:',
    '- Kiến trúc module-based: NestJS tổ chức code theo modules, mỗi module đóng gói controller + service + DTOs, phù hợp với hệ thống có nhiều domain nghiệp vụ (Auth, POI, Tour, ShopOwner, Tourist, TTS, QR, Analytics).',
    '- Dependency Injection: NestJS cung cấp DI container tích hợp, giúp quản lý dependencies giữa các service dễ dàng (ví dụ: PoisService inject TtsService, QrService, SeedExportService).',
    '- TypeScript first-class: NestJS được viết bằng TypeScript, hỗ trợ decorators (@Injectable, @Controller, @Roles, @UseGuards) giúp code clean và declarative.',
    '- Ecosystem: Tích hợp sẵn Swagger, Passport.js, class-validator, ServeStatic. Express đòi hỏi cấu hình thủ công cho mỗi thành phần.',
    '- Trade-off: NestJS có learning curve cao hơn Express và bundle size lớn hơn, nhưng phù hợp cho dự án trung bình-lớn cần cấu trúc rõ ràng.',
])

add_content('2. Prisma vs TypeORM', [
    'Nhóm chọn Prisma ORM thay vì TypeORM vì:',
    '- Schema-first approach: Prisma sử dụng schema.prisma file định nghĩa model, tự sinh TypeScript types. Đảm bảo DB schema và code luôn đồng bộ.',
    '- Type safety: Prisma Client tự sinh query types chính xác, IDE autocomplete tốt. TypeORM dùng decorator-based entity có thể bị mismatch.',
    '- Migration: Prisma Migrate tạo migration tự động từ schema changes, dễ quản lý version.',
    '- Prisma Studio: GUI web để browse/edit data trực tiếp, rất hữu ích trong development.',
    '- Performance: Prisma query engine tối ưu hơn TypeORM cho các truy vấn phức tạp với nested relations.',
    '- Trade-off: Prisma không hỗ trợ raw SQL linh hoạt bằng TypeORM, và không hỗ trợ eager/lazy loading truyền thống.',
])

add_content('3. Expo vs Flutter', [
    'Nhóm chọn Expo (React Native) thay vì Flutter vì:',
    '- JavaScript/TypeScript ecosystem: Toàn bộ tech stack (Backend NestJS, Frontend React, Mobile Expo) đều dùng TypeScript, giảm context switching và tận dụng chung kiến thức.',
    '- Expo Go: Cho phép test trực tiếp trên thiết bị mà không cần build native (quét QR code để mở app). Rất tiện cho phát triển nhanh.',
    '- Expo SDK: Cung cấp sẵn các module native (expo-location, expo-camera, expo-audio, expo-network) không cần cấu hình native code.',
    '- OTA updates: Expo hỗ trợ cập nhật over-the-air mà không cần re-submit app store.',
    '- Trade-off: Flutter có performance tốt hơn cho animation phức tạp và UI pixel-perfect. Tuy nhiên, GPS Tours chủ yếu dùng bản đồ + audio nên React Native đáp ứng đủ.',
])

add_content('4. PostgreSQL vs MongoDB', [
    'Nhóm chọn PostgreSQL thay vì MongoDB vì:',
    '- Dữ liệu có cấu trúc quan hệ rõ ràng: User ↔ ShopOwner (1:1), User ↔ POI (1:N), Tour ↔ POI (N:N qua TourPoi). PostgreSQL với referential integrity đảm bảo tính toàn vẹn.',
    '- ACID transactions: Quan trọng cho các thao tác như tạo Tour + TourPois trong 1 transaction.',
    '- Prisma support: Prisma hỗ trợ PostgreSQL tốt nhất với đầy đủ tính năng (migration, relation queries, aggregation).',
    '- Geospatial: PostgreSQL hỗ trợ PostGIS extension cho truy vấn địa lý (dù hiện tại nhóm dùng Haversine ở application layer).',
    '- Trade-off: MongoDB linh hoạt hơn cho dữ liệu phi cấu trúc, nhưng data model của GPS Tours có schema cố định nên PostgreSQL phù hợp hơn.',
])

add_content('1. Mối quan hệ giữa các module backend', [
    'Sơ đồ thành phần backend (NestJS modules) và dependencies:',
    '- AppModule (root) imports tất cả modules con.',
    '- PrismaModule (global): Cung cấp PrismaService cho mọi module.',
    '- ConfigModule (global): Cung cấp biến môi trường.',
    '- AuthModule: Sử dụng PrismaModule, JwtModule, PassportModule. Export JwtAuthGuard.',
    '- PoisModule: Import TtsModule, QrModule. Inject SeedExportService.',
    '- ToursModule: Inject SeedExportService.',
    '- ShopOwnerModule: Import PoisModule (reuse PoisService).',
    '- TtsModule: Sử dụng PrismaModule, msedge-tts library.',
    '- QrModule: Sử dụng PrismaModule, qrcode library.',
    '- MediaModule: Sử dụng Multer cho file upload.',
    '- PublicModule: Import PoisModule, ToursModule. Không require auth.',
    '- SeedExportModule (global): Export SeedExportService để PoisModule và ToursModule inject.',
    '(Sơ đồ chi tiết xem tại docs/step3_prd/17_component_diagram.md)',
])

add_content('2. Mối quan hệ giữa backend và các client', [
    'Sơ đồ thành phần giữa backend API và 3 client:',
    '- Admin Dashboard (React) → API: Sử dụng nhóm endpoint /auth/*, /pois/*, /tours/*, /admin/analytics/*, /merchants/*, /me/*. Gửi JWT Bearer token trên mọi request.',
    '- Shop Owner Dashboard (React, chung codebase) → API: Sử dụng /auth/*, /shop-owner/*, /tts/*. JWT token chứa role: SHOP_OWNER.',
    '- Tourist App (Expo) → API: Sử dụng /public/* (không cần auth) cho bản đồ, POI, Tour, QR validate. Sử dụng /auth/*, /tourist/* (cần auth) cho profile, favorites, history.',
    '- Tất cả client giao tiếp qua HTTP/HTTPS RESTful API, format JSON.',
    '- File media được serve trực tiếp qua /uploads/* (ServeStaticModule).',
])

print('Section B done.')

# ═══════════════════════════════════════════════════════════
# Save after A+B to keep progress
# ═══════════════════════════════════════════════════════════
doc.save(DST)
print(f'\nSaved to {DST}')
print('Sections A and B complete. Run fill_report_part2.py for C-F.')
